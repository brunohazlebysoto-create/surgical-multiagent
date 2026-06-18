import pytest
from docx import Document
from app.services.docx_generator import parse_inline_formatting

def test_parse_inline_formatting_no_bold():
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "Hello world")
    assert len(p.runs) == 1
    assert p.runs[0].text == "Hello world"
    assert p.runs[0].bold is None

def test_parse_inline_formatting_only_bold():
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "**Hello world**")
    assert len(p.runs) == 3
    assert p.runs[0].text == ""
    assert p.runs[1].text == "Hello world"
    assert p.runs[1].bold is True
    assert p.runs[2].text == ""

def test_parse_inline_formatting_mixed():
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "Normal **bold** normal")
    assert len(p.runs) == 3
    assert p.runs[0].text == "Normal "
    assert p.runs[1].text == "bold"
    assert p.runs[1].bold is True
    assert p.runs[2].text == " normal"

def test_parse_inline_formatting_multiple_bold():
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "**One** and **Two**")
    assert len(p.runs) == 5
    assert p.runs[0].text == ""
    assert p.runs[1].text == "One"
    assert p.runs[1].bold is True
    assert p.runs[2].text == " and "
    assert p.runs[3].text == "Two"
    assert p.runs[3].bold is True
    assert p.runs[4].text == ""

def test_parse_inline_formatting_empty_string():
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "")
    assert len(p.runs) == 1
    assert p.runs[0].text == ""

def test_parse_inline_formatting_empty_bold():
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "****")
    assert len(p.runs) == 3
    assert p.runs[0].text == ""
    assert p.runs[1].text == ""
    assert p.runs[1].bold is True
    assert p.runs[2].text == ""

def test_parse_inline_formatting_unmatched_asterisks():
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "**Hello world")
    assert len(p.runs) == 1
    assert p.runs[0].text == "**Hello world"
    assert p.runs[0].bold is None

def test_parse_inline_formatting_single_asterisks():
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "*Hello world*")
    assert len(p.runs) == 1
    assert p.runs[0].text == "*Hello world*"
    assert p.runs[0].bold is None

def test_parse_inline_formatting_three_asterisks_logic():
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "***Hello world***")
    # re.split gives ['', '***Hello world**', '*']
    assert len(p.runs) == 3
    assert p.runs[0].text == ""
    assert p.runs[1].text == "*Hello world"
    assert p.runs[1].bold is True
    assert p.runs[2].text == "*"
    assert p.runs[2].bold is None

def test_parse_inline_formatting_newlines():
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "**Hello\nworld**")
    # The current code uses r'(\*\*.*?\*\*)' which DOES NOT match across newlines without re.DOTALL
    # So the text is not split, and it's treated as a single normal run
    assert len(p.runs) == 1
    assert p.runs[0].text == "Hello\nworld"
    assert p.runs[0].bold is True


def test_parse_inline_formatting_empty_paragraph():
    doc = Document()
    p = doc.add_paragraph()
    parse_inline_formatting(p, "")
    assert len(p.runs) == 1
    assert p.runs[0].text == ""
    assert p.runs[0].bold is None
